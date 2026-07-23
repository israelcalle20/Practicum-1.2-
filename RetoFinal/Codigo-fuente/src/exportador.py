"""
exportador.py — ensambla el documento Word final a partir de los recursos
ya generados (leídos desde MongoDB por pipeline.py y pasados aquí como
lista de diccionarios).

Estructura del documento:
- Portada: nombre de la asignatura y fecha de generación.
- Una sección (heading nivel 1) por unidad.
- Cuatro subsecciones (heading nivel 2) por unidad: resumen, glosario,
  preguntas de comprensión y actividad de reflexión.
"""
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AZUL_UTPL = RGBColor(0x0B, 0x41, 0x6D)


def _agregar_portada(doc: DocxDocument, asignatura: str):
    titulo = doc.add_heading(level=0)
    run = titulo.add_run(f"Recursos didácticos — {asignatura}")
    run.font.color.rgb = AZUL_UTPL
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_run = fecha.add_run(f"Generado automáticamente — {datetime.now().strftime('%d/%m/%Y')}")
    fecha_run.italic = True
    doc.add_page_break()


def _agregar_resumen(doc: DocxDocument, resumen: dict):
    doc.add_heading("Resumen ejecutivo", level=2)
    if not resumen:
        doc.add_paragraph("(Recurso no disponible)")
        return
    doc.add_paragraph(resumen["parrafo_1"])
    doc.add_paragraph(resumen["parrafo_2"])


def _agregar_glosario(doc: DocxDocument, glosario: dict):
    doc.add_heading("Glosario de términos", level=2)
    if not glosario:
        doc.add_paragraph("(Recurso no disponible)")
        return
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].text = "Término"
    hdr[1].text = "Definición"
    for celda in hdr:
        for p in celda.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for termino in glosario["terminos"]:
        fila = tabla.add_row().cells
        fila[0].text = termino["termino"]
        fila[1].text = termino["definicion"]


def _agregar_preguntas(doc: DocxDocument, preguntas: dict):
    doc.add_heading("Preguntas de comprensión", level=2)
    if not preguntas:
        doc.add_paragraph("(Recurso no disponible)")
        return
    etiquetas_nivel = {"facil": "Fácil", "media": "Media", "dificil": "Difícil"}
    for i, item in enumerate(preguntas["preguntas"], start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. [{etiquetas_nivel.get(item['nivel'], item['nivel'])}] ").bold = True
        p.add_run(item["pregunta"])
        respuesta = doc.add_paragraph()
        respuesta.add_run("Respuesta esperada: ").italic = True
        respuesta.add_run(item["respuesta"])


def _agregar_actividad(doc: DocxDocument, actividad: dict):
    doc.add_heading("Actividad de reflexión", level=2)
    if not actividad:
        doc.add_paragraph("(Recurso no disponible)")
        return
    doc.add_paragraph(actividad["enunciado"])
    producto = doc.add_paragraph()
    producto.add_run("Producto esperado: ").bold = True
    producto.add_run(actividad["producto_esperado"])


def exportar_word(asignatura: str, unidades_recursos: list, ruta_salida: str):
    """
    unidades_recursos: lista de dicts con las llaves
    'numero', 'titulo', 'resumen', 'glosario', 'preguntas', 'actividad'
    (cada una ya en el formato JSON que devuelve prompts.py).
    """
    doc = DocxDocument()
    _agregar_portada(doc, asignatura)

    for ur in unidades_recursos:
        doc.add_heading(f"Unidad {ur['numero']}: {ur['titulo']}", level=1)
        _agregar_resumen(doc, ur.get("resumen"))
        _agregar_glosario(doc, ur.get("glosario"))
        _agregar_preguntas(doc, ur.get("preguntas"))
        _agregar_actividad(doc, ur.get("actividad"))
        doc.add_page_break()

    doc.save(ruta_salida)
    print(f"Documento guardado: {ruta_salida}")
