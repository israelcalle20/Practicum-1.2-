"""
extractor.py — extrae texto plano de los documentos fuente (PDF y Word),
y metadatos estructurados de la tabla de datos básicos del Plan Docente.

No decide todavía a qué unidad pertenece cada fragmento; eso lo hace
base_conocimiento.py / construir_base_real.py, que son quienes conocen
la estructura de la asignatura.
"""
from pathlib import Path
from pypdf import PdfReader
import docx


def extraer_texto_pdf(ruta: str) -> str:
    """Extrae todo el texto de un PDF, página por página."""
    lector = PdfReader(ruta)
    paginas = [pagina.extract_text() or "" for pagina in lector.pages]
    return "\n".join(paginas).strip()


def extraer_texto_docx(ruta: str) -> str:
    """Extrae todo el texto (párrafos) de un documento Word, sin tablas."""
    documento = docx.Document(ruta)
    parrafos = [p.text for p in documento.paragraphs if p.text.strip()]
    return "\n".join(parrafos).strip()


def extraer_texto(ruta: str) -> str:
    """Detecta el tipo de archivo por extensión y despacha al extractor correcto."""
    extension = Path(ruta).suffix.lower()
    if extension == ".pdf":
        return extraer_texto_pdf(ruta)
    elif extension in (".docx", ".doc"):
        return extraer_texto_docx(ruta)
    else:
        raise ValueError(f"Tipo de archivo no soportado: {extension}")


def extraer_metadatos_plan_docente(ruta_docx: str) -> dict:
    """
    Busca en TODAS las tablas del Plan Docente (no solo la primera) los
    campos de asignatura/código/carrera/ciclo, identificándolos por
    palabra clave en la etiqueta de cada fila — no por una posición o
    estructura de tabla fija. Esto es necesario porque distintas
    plantillas de la UTPL organizan estos datos de forma distinta:
    algunas usan una fila por dato ("Codigo | COMP_2010"), otras usan
    etiquetas más largas con celdas combinadas y repetidas
    ("Código de la asignatura | DSOF_2036 | DSOF_2036 | DSOF_2036 | ...").

    Devuelve un diccionario con llaves en minúsculas y sin tildes:
    'asignatura', 'codigo', 'carrera', 'ciclo' (y, si se encuentran,
    'periodo' y 'prerequisito').
    """
    documento = docx.Document(ruta_docx)
    if not documento.tables:
        raise ValueError(
            f"'{ruta_docx}' no tiene tablas — no se pudo extraer los datos "
            "básicos. Revisa que sea el Plan Docente correcto."
        )

    reemplazos_tildes = str.maketrans("áéíóúÁÉÍÓÚñÑ", "aeiouAEIOUnN")

    # Orden importa: se revisa 'codigo' antes que 'asignatura' porque una
    # etiqueta como "Código de la asignatura" contiene la palabra
    # "asignatura" y no debe confundirse con el nombre de la asignatura.
    patrones = {
        "codigo": ["codigo"],
        "asignatura": ["nombre de la asignatura", "asignatura"],
        "carrera": ["carrera"],
        "ciclo": ["nivel", "ciclo"],
        "periodo": ["periodo academico", "semestre"],
        "prerequisito": ["prerequisito", "prerrequisito"],
    }
    esenciales = ["codigo", "asignatura", "carrera", "ciclo"]

    metadatos = {}
    for tabla in documento.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells]
            if len(celdas) < 2 or not celdas[0]:
                continue
            etiqueta = celdas[0].translate(reemplazos_tildes).lower()
            # El valor real es la primera celda distinta de la etiqueta
            # (las celdas combinadas de Word aparecen repetidas varias
            # veces con el mismo texto al leerlas con python-docx).
            valor = next((c for c in celdas[1:] if c and c != celdas[0]), "")
            if not valor:
                continue
            for clave, palabras_clave in patrones.items():
                if clave in metadatos:
                    continue
                if any(palabra in etiqueta for palabra in palabras_clave):
                    metadatos[clave] = valor
                    break
        if all(clave in metadatos for clave in esenciales):
            break  # ya tenemos lo esencial, no hace falta seguir escaneando

    faltantes = [clave for clave in esenciales if clave not in metadatos]
    if faltantes:
        raise ValueError(
            f"No se pudieron detectar estos datos en '{ruta_docx}': {', '.join(faltantes)}. "
            "Revisa que el documento tenga una tabla de Datos Básicos con esas etiquetas, "
            "o pásalos manualmente."
        )

    return metadatos


if __name__ == "__main__":
    # Prueba rápida manual: ajusta la ruta a un archivo real de tu asignatura
    texto = extraer_texto("data/plan_docente.pdf")
    print(texto[:500])